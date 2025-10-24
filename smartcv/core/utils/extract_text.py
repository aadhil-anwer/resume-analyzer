
from core.utils.normalize import normalize_text
import fitz  # PyMuPDF
import docx
import pytesseract
pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
from PIL import Image
# ------------------ TEXT EXTRACTION ------------------
def extract_text_from_pdf(path: str) -> str:
    """
    Fully resilient PDF extractor:
    1. Normal text extraction (PyMuPDF)
    2. Fallback to block extraction for multi-column layouts
    3. OCR fallback for image/scanned pages
    4. Never returns empty string
    """
    text_content = []
    try:
        with fitz.open(path) as pdf:
            for page_index, page in enumerate(pdf):
                page_text = page.get_text("text")

                # Block extraction if text is empty or too short
                if not page_text.strip() or len(page_text.strip()) < 40:
                    blocks = page.get_text("blocks")
                    page_text = "\n".join(
                        b[4] for b in blocks if b[4].strip()
                    )

                # OCR fallback if still blank
                if not page_text.strip():
                    pix = page.get_pixmap(dpi=300)
                    img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                    ocr_text = pytesseract.image_to_string(img)
                    page_text = ocr_text.strip() or f"[OCR failed on page {page_index + 1}]"

                text_content.append(page_text)
    except Exception as e:
        return f"[ERROR extracting PDF text: {e}]"

    combined = "\n".join(text_content)
    print(combined)
    return normalize_text(combined)



def extract_text_from_docx(path: str) -> str:
    """
    Robust DOCX extractor:
    - Reads paragraphs, tables, headers/footers
    - Gracefully handles malformed documents
    """
    text_content = []
    try:
        doc = docx.Document(path)

        # Paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_content.append(para.text.strip())

        # Tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_content.append(row_text)

        # Headers and footers (optional)
        for section in doc.sections:
            if hasattr(section, "header"):
                header_text = " ".join(
                    p.text.strip() for p in section.header.paragraphs if p.text.strip()
                )
                if header_text:
                    text_content.append(header_text)
            if hasattr(section, "footer"):
                footer_text = " ".join(
                    p.text.strip() for p in section.footer.paragraphs if p.text.strip()
                )
                if footer_text:
                    text_content.append(footer_text)
    except Exception as e:
        return f"[ERROR extracting DOCX text: {e}]"

    combined = "\n".join(text_content)
    return normalize_text(combined)